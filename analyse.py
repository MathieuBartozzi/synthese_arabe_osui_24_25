# === NOUVELLE VERSION DU SCRIPT : synthese_reseau_OSUI.py ===
# Objectif : Synthèse exhaustive par établissement, puis synthèse réseau à partir de celles-ci

import os, requests, pandas as pd
from tqdm import tqdm
from docx import Document
from collections import defaultdict, Counter
from typing import List
from pydantic import BaseModel
from openai import OpenAI
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from dotenv import load_dotenv
import os

# Charger les variables d'environnement
load_dotenv()

# Récupérer la clé
api_key = os.getenv("OPENAI_API_KEY")

# Utilisation avec OpenAI
client = OpenAI(api_key=api_key)

# === CONFIGURATION ===
INPUT_CSV = "rapports.csv"
RAW_DIR = "raw_docs"
OUTPUT_DOC = "rapport_final_OSUI.docx"
os.makedirs(RAW_DIR, exist_ok=True)

# === MODELES ===
class RapportClasse(BaseModel):
    points_forts: str
    difficultes_identifiees: str
    pratiques_interet: str
    recommandations_deduites: str

class SyntheseEtab(BaseModel):
    synthese: str

class TableauTheme(BaseModel):
    theme: str
    constat: str
    reco: str
    acteurs: str

class TableauPilotage(BaseModel):
    tableau: List[TableauTheme]

# === OUTILS ===
def get_download_url(view_url):
    if "document/d/" in view_url:
        file_id = view_url.split("/document/d/")[1].split("/")[0]
        return f"https://docs.google.com/document/d/{file_id}/export?format=docx"
    return None

def download_docx(url, path):
    r = requests.get(url)
    if r.status_code == 200:
        with open(path, "wb") as f: f.write(r.content)
        return True
    return False

def extract_text(path):
    try:
        return "\n".join([p.text for p in Document(path).paragraphs])
    except:
        return ""

def detecter_doublons(df, seuil=0.95):
    textes = df['texte'].fillna("").tolist()
    vect = TfidfVectorizer().fit_transform(textes)
    sim = cosine_similarity(vect)
    to_drop = set()
    for i in range(len(sim)):
        for j in range(i + 1, len(sim)):
            if sim[i][j] > seuil:
                to_drop.add(j)
    df_cleaned = df.drop(index=df.index[list(to_drop)]).reset_index(drop=True)
    print(f"🧹 {len(to_drop)} doublons supprimés.")
    return df_cleaned

def analyser_rapport(texte):
    try:
        completion = client.beta.chat.completions.parse(
            model="gpt-4o",
            messages=[{
                "role": "system",
                "content": (
                    "Tu es conseiller pédagogique spécialisé. Analyse ce texte de rapport de visite de classe et produis :\n"
                    "- points_forts\n- difficultes_identifiees\n- pratiques_d'interet\n- recommandations_deduites"
                    "Reste strictement fidèle aux éléments fournis. N’ajoute rien qui ne soit présent explicitement dans les données. Aucune généralisation excessive ni extrapolation.\n"

                )
            }, {"role": "user", "content": texte}],
            response_format=RapportClasse
        )
        return completion.choices[0].message.parsed
    except Exception as e:
        print("Erreur analyse GPT :", e)
        return None

def synthese_par_etablissement(observations):
    try:
        prompt = (
            "Tu es conseiller pédagogique. Voici un ensemble d'observations issues de rapports de visite de classe pour un établissement et pur l'enseignmement de la langue Arabe.\n"
            "Rédige une synthèse structurée en plusieurs paragraphes, incluant :\n"
            "- Points forts collectifs\n- Difficultés pédagogiques récurrentes\n- pratiques d'interet observées\n- Recommandations pour l’équipe\n"
            "Reste strictement fidèle aux éléments fournis. N’ajoute rien qui ne soit présent explicitement dans les données. Aucune généralisation excessive ni extrapolation.\n"
            "Utilise un style professionnel, collectif et synthétique. Evite les redondances.\n"
            "N’utilise aucun formatage Markdown ou typographique (pas de #, **, ou _). Rédige en texte brut uniquement.\n\n"

        )
        texte = "\n\n".join([
            f"Points forts : {o['points_forts']}\nDifficultés : {o['difficultes_identifiees']}\nPratiques : {o['pratiques_interet']}\nRecommandations : {o['recommandations_deduites']}"
            for o in observations
        ])

        completion = client.beta.chat.completions.parse(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt + texte}],
            response_format=SyntheseEtab
        )
        return completion.choices[0].message.parsed.synthese
    except Exception as e:
        print("Erreur synthèse établissement :", e)
        return ""

def synthese_reseau_depuis_etablissements(syntheses_etab):
    try:
        message = (
            "Voici les synthèses pédagogiques de plusieurs établissements faisant partie d'un reseau.\n"
            "Ta tâche est de rédiger une synthèse pour le pilotage du réseau :\n"
            "- Identifier les convergences et spécificités\n"
            "- Extraire 3 à 5 enjeux majeurs\n"
            "- Rédiger une conclusion stratégique claire et exploitable pour la direction du reseau.\n"
            "Reste strictement fidèle aux éléments fournis. N’ajoute rien qui ne soit présent explicitement dans les données. Aucune généralisation excessive ni extrapolation.\n"
            "Utilise un style synthétique et professionnel.\n\n"
            "N’utilise aucun formatage Markdown ou typographique (pas de #, **, ou _). Rédige en texte brut uniquement.\n\n"

            + "\n\n".join(syntheses_etab)
        )
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": message}]
        )
        return completion.choices[0].message.content.strip()
    except Exception as e:
        print("Erreur synthèse réseau :", e)
        return ""

def add_section(doc, titre, contenu):
    doc.add_heading(titre, level=2)
    doc.add_paragraph(contenu)

# === PIPELINE PRINCIPAL ===
df = pd.read_csv(INPUT_CSV)
df['texte'] = None
for idx, row in tqdm(df.iterrows(), total=len(df)):
    path = f"{RAW_DIR}/{idx}.docx"
    url = get_download_url(row['lien'])
    if not url or not download_docx(url, path): continue
    texte = extract_text(path)
    df.at[idx, 'texte'] = texte

df = detecter_doublons(df)
observations = []
for idx, row in tqdm(df.iterrows(), total=len(df)):
    if pd.isna(row['texte']): continue
    analyse = analyser_rapport(row['texte'])
    if analyse:
        observations.append({"etablissement": row['etablissement'], **analyse.model_dump()})

# === SYNTHÈSES PAR ETABLISSEMENT ===
obs_par_etab = defaultdict(list)
for obs in observations:
    obs_par_etab[obs['etablissement']].append(obs)

synth_etab = {}
for etab, obs in obs_par_etab.items():
    synth_etab[etab] = synthese_par_etablissement(obs)

# === SYNTHÈSE RESEAU ===
synth_reseau = synthese_reseau_depuis_etablissements(list(synth_etab.values()))

# === DOCUMENT FINAL ===
doc = Document()
doc.add_heading("Synthèse pédagogique OSUI", 0)

# Partie 1 : Synthèse réseau
doc.add_heading("1. Synthèse réseau", level=1)
doc.add_paragraph(synth_reseau)

# Partie 2 : Synthèses par établissement
doc.add_heading("2. Synthèses par établissement", level=1)
for etab, contenu in synth_etab.items():
    add_section(doc, etab, contenu)

doc.save(OUTPUT_DOC)
print("✅ Rapport final généré.")
